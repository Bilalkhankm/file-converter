# ╔══════════════════════════════════════════════════════════════════════════╗
# ║           AI-POWERED FILE CONVERTER PRO — app.py                       ║
# ║           Powered by Groq LLaMA · Built with Gradio                    ║
# ║                                                                          ║
# ║  HOW TO RUN:                                                             ║
# ║    1. pip install -r requirements.txt                                    ║
# ║    2. Set env variable:  GROQ_API_KEY=your_key_here                      ║
# ║    3. python app.py                                                      ║
# ╚══════════════════════════════════════════════════════════════════════════╝

# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 1 — IMPORTS
# ═══════════════════════════════════════════════════════════════════════════

import os
import io
import re
import json
import base64
import tempfile
import time
import datetime
import csv as csv_mod
from pathlib import Path
from typing import Optional, List, Tuple

import pandas as pd
import chardet
import gradio as gr

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

from PIL import Image
from groq import Groq

# ── Optional libraries (graceful fallback if not installed) ───────────────

try:
    from pdf2docx import Converter as PDFConverter
    PDF2DOCX_AVAILABLE = True
except ImportError:
    PDF2DOCX_AVAILABLE = False
    print("⚠️  pdf2docx not found — standard PDF→Word disabled")

try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False
    print("⚠️  pdf2image not found — scanned PDF OCR disabled")

try:
    import pytesseract
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False
    print("⚠️  pytesseract not found — Tesseract OCR fallback disabled")


# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 2 — GROQ API SETUP
# ═══════════════════════════════════════════════════════════════════════════

# Load API key from environment variable
GROQ_API_KEY = os.environ.get("GROQ_API_KEY", "")

try:
    if not GROQ_API_KEY:
        raise ValueError("GROQ_API_KEY environment variable not set.")

    groq_client = Groq(api_key=GROQ_API_KEY)

    # Quick connectivity test
    _test = groq_client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": "Say OK"}],
        max_tokens=5
    )
    GROQ_AVAILABLE = True
    print(f"✅ Groq connected! Model: llama-3.3-70b-versatile")
    print(f"   Test response: {_test.choices[0].message.content}")

except Exception as e:
    GROQ_AVAILABLE = False
    groq_client   = None
    print(f"⚠️  Groq connection failed: {e}")
    print("   App will run in Standard Mode (no AI features).")


# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 3 — GROQ AI HELPER FUNCTIONS
# ═══════════════════════════════════════════════════════════════════════════

# ── Helper 1: Image/scanned page se text nikalo ──────────────────────────
def groq_extract_text_from_image(image_path: str) -> str:
    """
    Use Groq Vision to extract text from a scanned PDF page image.
    Falls back to Tesseract OCR if Groq is unavailable.
    """
    if not GROQ_AVAILABLE:
        if TESSERACT_AVAILABLE:
            return pytesseract.image_to_string(Image.open(image_path))
        return "[OCR unavailable — no Groq API key or Tesseract installed]"

    try:
        with open(image_path, "rb") as f:
            b64_image = base64.b64encode(f.read()).decode("utf-8")

        ext  = Path(image_path).suffix.lower()
        mime = {
            ".png": "image/png",
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".webp": "image/webp"
        }.get(ext, "image/png")

        response = groq_client.chat.completions.create(
            model="llama-3.2-90b-vision-preview",
            messages=[{
                "role": "user",
                "content": [
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:{mime};base64,{b64_image}"}
                    },
                    {
                        "type": "text",
                        "text": (
                            "Extract ALL text from this document image with high accuracy. "
                            "Preserve the original structure: headings, paragraphs, tables, "
                            "lists, bullet points, and numbering. "
                            "Use markdown-style formatting: # for headings, - for bullets, "
                            "| for table cells. Do not add commentary — return only the text."
                        )
                    }
                ]
            }],
            max_tokens=4096
        )
        return response.choices[0].message.content

    except Exception as e:
        if TESSERACT_AVAILABLE:
            print(f"⚠️  Groq Vision failed ({e}), falling back to Tesseract...")
            return pytesseract.image_to_string(Image.open(image_path))
        return f"[Text extraction failed: {e}]"


# ── Helper 2: Text se tables dhundo aur nikalo ───────────────────────────
def groq_extract_tables_from_text(text: str) -> List[dict]:
    """
    Use Groq LLM to identify and extract table-like structures
    from unstructured document text.
    Returns list of dicts: [{name, headers, rows}, ...]
    """
    if not GROQ_AVAILABLE:
        return []

    prompt = f"""
Analyze the following document text and extract ALL structured data:
- Tables (explicit or implicit)
- Lists of items with attributes
- Key-value pairs (like invoice fields)
- Grade tables, product lists, schedules, etc.

Return a JSON array where each element is a table object:
{{
  "name": "<descriptive table name>",
  "headers": ["col1", "col2", ...],
  "rows": [["val1", "val2", ...], ...]
}}

Return ONLY valid JSON. No explanation or extra text.

Document text:
---
{text[:6000]}
---
"""
    try:
        response = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=4096,
            temperature=0.1
        )
        raw = response.choices[0].message.content.strip()
        # Strip markdown code fences if present
        raw = re.sub(r'^```json\s*|^```\s*|\s*```$', '', raw, flags=re.MULTILINE).strip()
        tables = json.loads(raw)
        return tables if isinstance(tables, list) else []

    except Exception as e:
        print(f"⚠️  Groq table extraction failed: {e}")
        return []


# ── Helper 3: Document summarization ────────────────────────────────────
def groq_summarize_document(text: str, style: str = "concise") -> str:
    """
    Summarize a document using Groq LLaMA.
    style options: 'concise' | 'detailed' | 'bullet'
    """
    if not GROQ_AVAILABLE:
        return "[Groq API not available — add GROQ_API_KEY to environment variables]"

    style_prompts = {
        "concise":  "Provide a concise 2-3 sentence summary.",
        "detailed": "Provide a detailed summary covering all main points.",
        "bullet":   "Provide a bullet-point summary of key facts and figures."
    }

    try:
        response = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{
                "role": "user",
                "content": (
                    f"{style_prompts.get(style, style_prompts['concise'])}\n\n"
                    f"Document:\n{text[:8000]}"
                )
            }],
            max_tokens=1024
        )
        return response.choices[0].message.content

    except Exception as e:
        return f"[Summarization failed: {e}]"


# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 4 — UTILITY HELPERS
# ═══════════════════════════════════════════════════════════════════════════

def make_output_path(input_path: str, new_ext: str) -> str:
    """Generate a safe output file path in a temp directory."""
    stem = Path(input_path).stem
    tmp  = tempfile.mkdtemp()
    return os.path.join(tmp, f"{stem}_converted{new_ext}")


def read_docx_text(docx_path: str) -> str:
    """Extract all paragraph text from a Word document."""
    doc = Document(docx_path)
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])


def style_excel_sheet(ws):
    """Apply professional navy/white styling to an Excel worksheet."""
    header_fill = PatternFill(start_color="1F3864", end_color="1F3864", fill_type="solid")
    alt_fill    = PatternFill(start_color="EBF0FA", end_color="EBF0FA", fill_type="solid")
    border_side = Side(style="thin", color="CCCCCC")
    thin_border = Border(
        left=border_side, right=border_side,
        top=border_side, bottom=border_side
    )

    # Style header row (row 1)
    for cell in ws[1]:
        cell.font      = Font(bold=True, color="FFFFFF", size=11)
        cell.fill      = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border    = thin_border

    # Style data rows (row 2 onwards)
    for i, row in enumerate(ws.iter_rows(min_row=2), start=2):
        fill = alt_fill if i % 2 == 0 else None
        for cell in row:
            if fill:
                cell.fill  = fill
            cell.alignment = Alignment(vertical="center", wrap_text=True)
            cell.border    = thin_border

    # Auto-fit column widths
    for col in ws.columns:
        max_len = max((len(str(c.value or "")) for c in col), default=10)
        ws.column_dimensions[get_column_letter(col[0].column)].width = min(max_len + 4, 50)

    ws.row_dimensions[1].height = 30


# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 5 — CONVERSION FUNCTIONS  [1/7 to 7/7]
# ═══════════════════════════════════════════════════════════════════════════

# ── [1/7] PDF → WORD ──────────────────────────────────────────────────────
def convert_pdf_to_word(pdf_path: str, use_groq: bool = False) -> Tuple[str, str]:
    """
    Convert PDF to Word (.docx).
    - Standard mode: uses pdf2docx (fast, for text-based PDFs)
    - Groq AI mode:  converts pages to images → Groq Vision OCR (for scanned PDFs)
    Returns (output_path, status_message)
    """
    if not pdf_path or not os.path.exists(pdf_path):
        return None, "❌ No valid PDF file provided."

    out_path = make_output_path(pdf_path, ".docx")

    # Standard mode — pdf2docx
    if not use_groq and PDF2DOCX_AVAILABLE:
        try:
            cv = PDFConverter(pdf_path)
            cv.convert(out_path, start=0, end=None)
            cv.close()
            return out_path, "✅ PDF converted to Word using pdf2docx (standard mode)."
        except Exception as e:
            print(f"⚠️  pdf2docx failed ({e}), trying OCR fallback...")

    # Groq AI / OCR mode — pdf2image + Vision
    if PDF2IMAGE_AVAILABLE:
        try:
            with tempfile.TemporaryDirectory() as tmp_dir:
                pages = convert_from_path(pdf_path, dpi=200)
                doc   = Document()

                # Base document styling
                normal_style           = doc.styles["Normal"]
                normal_style.font.name = "Calibri"
                normal_style.font.size = Pt(11)

                for page_num, page_img in enumerate(pages, 1):
                    img_path = os.path.join(tmp_dir, f"page_{page_num}.png")
                    page_img.save(img_path, "PNG")

                    if page_num > 1:
                        doc.add_page_break()

                    # Page heading
                    h = doc.add_heading(f"Page {page_num}", level=2)
                    if h.runs:
                        h.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

                    # Extract text via Groq Vision or Tesseract
                    extracted = groq_extract_text_from_image(img_path)

                    # Parse markdown-ish output into Word formatting
                    for line in extracted.split("\n"):
                        line = line.rstrip()
                        if not line:
                            doc.add_paragraph()
                        elif line.startswith("# "):
                            doc.add_heading(line[2:], level=1)
                        elif line.startswith("## "):
                            doc.add_heading(line[3:], level=2)
                        elif line.startswith("### "):
                            doc.add_heading(line[4:], level=3)
                        elif line.startswith(("- ", "* ", "• ")):
                            doc.add_paragraph(line[2:].strip(), style="List Bullet")
                        elif re.match(r"^\d+\.\s", line):
                            doc.add_paragraph(line, style="List Number")
                        else:
                            doc.add_paragraph(line)

                doc.save(out_path)
                mode = "Groq AI Vision" if (use_groq and GROQ_AVAILABLE) else "Tesseract OCR"
                return out_path, f"✅ Scanned PDF converted using {mode} ({len(pages)} pages)."

        except Exception as e:
            return None, f"❌ Conversion failed: {e}"

    return None, "❌ No conversion method available. Install pdf2image + poppler-utils."


# ── [2/7] WORD → PDF ──────────────────────────────────────────────────────
def convert_word_to_pdf(docx_path: str) -> Tuple[str, str]:
    """
    Convert Word (.docx) to PDF using LibreOffice headless.
    Falls back to docx2pdf if LibreOffice is unavailable.
    Returns (output_path, status_message)
    """
    if not docx_path or not os.path.exists(docx_path):
        return None, "❌ No valid Word file provided."

    try:
        out_dir = tempfile.mkdtemp()
        result  = os.popen(
            f'libreoffice --headless --convert-to pdf "{docx_path}" --outdir "{out_dir}" 2>&1'
        ).read()

        stem    = Path(docx_path).stem
        out_pdf = os.path.join(out_dir, f"{stem}.pdf")

        if os.path.exists(out_pdf):
            return out_pdf, "✅ Word converted to PDF using LibreOffice."

        # Fallback: docx2pdf
        try:
            from docx2pdf import convert as d2p
            out_pdf2 = make_output_path(docx_path, ".pdf")
            d2p(docx_path, out_pdf2)
            return out_pdf2, "✅ Word converted to PDF using docx2pdf."
        except Exception as e2:
            return None, (
                f"❌ PDF conversion failed.\n"
                f"LibreOffice output: {result}\n"
                f"Fallback error: {e2}"
            )

    except Exception as e:
        return None, f"❌ Conversion failed: {e}"


# ── [3/7] WORD → EXCEL ────────────────────────────────────────────────────
def convert_word_to_excel(docx_path: str, use_groq: bool = True) -> Tuple[str, str]:
    """
    Extract structured data from a Word document into Excel.
    - Always extracts native Word tables
    - With Groq AI: also extracts implicit tables from narrative text
    Returns (output_path, status_message)
    """
    if not docx_path or not os.path.exists(docx_path):
        return None, "❌ No valid Word file provided."

    out_path     = make_output_path(docx_path, ".xlsx")
    doc          = Document(docx_path)
    wb           = openpyxl.Workbook()
    tables_added = 0

    # Extract native Word tables
    for i, table in enumerate(doc.tables, 1):
        ws = wb.create_sheet(title=f"Table_{i}")
        for r_idx, row in enumerate(table.rows, 1):
            for c_idx, cell in enumerate(row.cells, 1):
                ws.cell(row=r_idx, column=c_idx, value=cell.text.strip())
        style_excel_sheet(ws)
        tables_added += 1

    # Groq AI: extract implicit/narrative tables
    if use_groq and GROQ_AVAILABLE:
        full_text = read_docx_text(docx_path)
        ai_tables = groq_extract_tables_from_text(full_text)

        for tbl in ai_tables:
            name    = str(tbl.get("name", f"AI_Table_{tables_added + 1}"))[:31]
            headers = tbl.get("headers", [])
            rows    = tbl.get("rows", [])

            if not headers and not rows:
                continue

            # Avoid duplicate sheet names
            sheet_name = name
            count      = 1
            while sheet_name in [s.title for s in wb.worksheets]:
                sheet_name = f"{name[:28]}_{count}"
                count += 1

            ws = wb.create_sheet(title=sheet_name)
            if headers:
                ws.append(headers)
            for row in rows:
                ws.append([str(v) for v in row])
            style_excel_sheet(ws)
            tables_added += 1

    # Remove default empty sheet if other sheets were created
    if tables_added > 0 and "Sheet" in wb.sheetnames:
        del wb["Sheet"]

    # Fallback: dump raw paragraphs if no tables found
    if tables_added == 0:
        ws       = wb.active
        ws.title = "Document_Text"
        ws.append(["Paragraph #", "Content"])
        for i, para in enumerate(doc.paragraphs, 1):
            if para.text.strip():
                ws.append([i, para.text.strip()])
        style_excel_sheet(ws)

    wb.save(out_path)
    mode = "Groq AI + " if (use_groq and GROQ_AVAILABLE) else ""
    return out_path, f"✅ Word converted to Excel ({mode}{tables_added} sheet(s) created)."


# ── [4/7] EXCEL → WORD ────────────────────────────────────────────────────
def convert_excel_to_word(xlsx_path: str) -> Tuple[str, str]:
    """
    Convert Excel workbook into a formatted Word report.
    Each sheet becomes a styled table in the Word document.
    Returns (output_path, status_message)
    """
    if not xlsx_path or not os.path.exists(xlsx_path):
        return None, "❌ No valid Excel file provided."

    out_path = make_output_path(xlsx_path, ".docx")

    try:
        wb  = openpyxl.load_workbook(xlsx_path)
        doc = Document()

        # Document title
        title_para           = doc.add_heading(f"Report: {Path(xlsx_path).stem}", 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Generation date
        date_para           = doc.add_paragraph(
            f"Generated: {datetime.datetime.now().strftime('%B %d, %Y at %H:%M')}"
        )
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            df = pd.DataFrame(ws.values)

            if df.empty:
                continue

            doc.add_heading(sheet_name, level=1)

            # Use first row as column headers
            df.columns = df.iloc[0]
            df         = df[1:].reset_index(drop=True)

            if df.empty:
                continue

            # Create styled Word table
            tbl = doc.add_table(
                rows=1 + len(df),
                cols=len(df.columns),
                style="Table Grid"
            )

            # Header row — navy background, white bold text
            for j, col_name in enumerate(df.columns):
                cell      = tbl.rows[0].cells[j]
                cell.text = str(col_name) if col_name else ""
                runs      = cell.paragraphs[0].runs
                if runs:
                    runs[0].bold             = True
                    runs[0].font.color.rgb   = RGBColor(0xFF, 0xFF, 0xFF)
                # Apply navy fill via XML
                tc  = cell._tc
                tcp = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "1F3864")
                shd.set(qn("w:val"),  "clear")
                tcp.append(shd)

            # Data rows
            for i, row_data in df.iterrows():
                for j, val in enumerate(row_data):
                    tbl.rows[i + 1].cells[j].text = str(val) if pd.notna(val) else ""

            doc.add_paragraph()

        doc.save(out_path)
        return out_path, f"✅ Excel converted to Word ({len(wb.sheetnames)} sheet(s) → tables)."

    except Exception as e:
        return None, f"❌ Excel → Word conversion failed: {e}"


# ── [5/7] CSV → EXCEL ─────────────────────────────────────────────────────
def convert_csv_to_excel(csv_path: str, delimiter: str = "auto") -> Tuple[str, str]:
    """
    Convert CSV file to styled Excel workbook.
    Auto-detects encoding and delimiter.
    Returns (output_path, status_message)
    """
    if not csv_path or not os.path.exists(csv_path):
        return None, "❌ No valid CSV file provided."

    out_path = make_output_path(csv_path, ".xlsx")

    try:
        # Detect file encoding
        with open(csv_path, "rb") as f:
            detected = chardet.detect(f.read())
        encoding = detected.get("encoding", "utf-8") or "utf-8"

        # Auto-detect delimiter
        if delimiter == "auto":
            with open(csv_path, "r", encoding=encoding, errors="replace") as f:
                sample = f.read(2048)
            try:
                sniffer   = csv_mod.Sniffer()
                delimiter = sniffer.sniff(sample).delimiter
            except Exception:
                delimiter = ","

        df = pd.read_csv(
            csv_path,
            encoding=encoding,
            sep=delimiter,
            encoding_errors="replace"
        )

        wb       = openpyxl.Workbook()
        ws       = wb.active
        ws.title = Path(csv_path).stem[:31]

        ws.append(list(df.columns))
        for _, row in df.iterrows():
            ws.append(list(row))

        style_excel_sheet(ws)
        wb.save(out_path)
        return out_path, (
            f"✅ CSV converted to Excel "
            f"({len(df):,} rows, {len(df.columns)} columns, encoding: {encoding})."
        )

    except Exception as e:
        return None, f"❌ CSV → Excel conversion failed: {e}"


# ── [6/7] EXCEL → CSV ─────────────────────────────────────────────────────
def convert_excel_to_csv(
    xlsx_path: str,
    sheet_name: str = "first",
    delimiter: str  = ","
) -> Tuple[str, str]:
    """
    Export an Excel sheet to CSV format.
    Returns (output_path, status_message)
    """
    if not xlsx_path or not os.path.exists(xlsx_path):
        return None, "❌ No valid Excel file provided."

    out_path = make_output_path(xlsx_path, ".csv")

    try:
        if sheet_name == "first":
            df = pd.read_excel(xlsx_path, engine="openpyxl")
        else:
            df = pd.read_excel(xlsx_path, sheet_name=sheet_name, engine="openpyxl")

        df.to_csv(out_path, index=False, sep=delimiter, encoding="utf-8-sig")
        return out_path, f"✅ Excel converted to CSV ({len(df):,} rows exported)."

    except Exception as e:
        return None, f"❌ Excel → CSV conversion failed: {e}"


# ── [7/7] BATCH CONVERSION ────────────────────────────────────────────────
def batch_convert(
    files: List[str],
    target_format: str,
    use_groq: bool = False
) -> Tuple[List[str], str]:
    """
    Convert multiple files to a single target format.
    Returns (list_of_output_paths, status_log_string)
    """
    if not files:
        return [], "❌ No files provided."

    CONVERTERS = {
        "docx": lambda f: convert_pdf_to_word(f, use_groq),
        "pdf":  lambda f: convert_word_to_pdf(f),
        "xlsx": lambda f: (
            convert_csv_to_excel(f)
            if f.lower().endswith(".csv")
            else convert_word_to_excel(f, use_groq)
        ),
        "csv":  lambda f: convert_excel_to_csv(f),
    }

    converter = CONVERTERS.get(target_format.lower().lstrip("."))
    if not converter:
        return [], f"❌ Unsupported target format: {target_format}"

    outputs = []
    log     = []

    for i, fp in enumerate(files, 1):
        name = Path(fp).name
        log.append(f"[{i}/{len(files)}] Processing: {name}")
        try:
            out, msg = converter(fp)
            if out:
                outputs.append(out)
            log.append(f"  → {msg}")
        except Exception as e:
            log.append(f"  → ❌ Error: {e}")

    log.append(f"\n✅ Batch complete: {len(outputs)}/{len(files)} files converted.")
    return outputs, "\n".join(log)


# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 6 — GRADIO HANDLER WRAPPERS
#  (Gradio UI ke liye thin wrappers — file.name path extract karte hain)
# ═══════════════════════════════════════════════════════════════════════════

def handle_pdf_to_word(file, use_groq):
    if file is None:
        return None, "⚠️ Please upload a PDF file."
    return convert_pdf_to_word(file.name, use_groq)


def handle_word_to_pdf(file):
    if file is None:
        return None, "⚠️ Please upload a Word (.docx) file."
    return convert_word_to_pdf(file.name)


def handle_word_to_excel(file, use_groq):
    if file is None:
        return None, "⚠️ Please upload a Word (.docx) file."
    return convert_word_to_excel(file.name, use_groq)


def handle_excel_to_word(file):
    if file is None:
        return None, "⚠️ Please upload an Excel (.xlsx) file."
    return convert_excel_to_word(file.name)


def handle_csv_to_excel(file, delimiter):
    if file is None:
        return None, "⚠️ Please upload a CSV file."
    delim_map = {
        "Auto-detect":  "auto",
        "Comma (,)":    ",",
        "Semicolon (;)": ";",
        "Tab (\\t)":     "\t",
        "Pipe (|)":      "|"
    }
    return convert_csv_to_excel(file.name, delim_map.get(delimiter, "auto"))


def handle_excel_to_csv(file, delimiter):
    if file is None:
        return None, "⚠️ Please upload an Excel (.xlsx) file."
    delim_map = {
        "Comma (,)":    ",",
        "Semicolon (;)": ";",
        "Tab (\\t)":     "\t",
        "Pipe (|)":      "|"
    }
    return convert_excel_to_csv(file.name, delimiter=delim_map.get(delimiter, ","))


def handle_summarize(file, style):
    if file is None:
        return "⚠️ Please upload a file to summarize."

    fp  = file.name
    ext = Path(fp).suffix.lower()

    try:
        if ext == ".docx":
            text = read_docx_text(fp)
        elif ext == ".txt":
            with open(fp, "r", errors="replace") as f:
                text = f.read()
        elif ext == ".csv":
            df   = pd.read_csv(fp, encoding_errors="replace")
            text = df.to_string()
        elif ext in (".xlsx", ".xls"):
            df   = pd.read_excel(fp, engine="openpyxl")
            text = df.to_string()
        else:
            return f"❌ Unsupported file format: {ext}"

        return groq_summarize_document(text, style.lower())

    except Exception as e:
        return f"❌ Summarization error: {e}"


def handle_batch(files, target_fmt, use_groq):
    if not files:
        return [], "⚠️ Please upload files to convert."
    paths       = [f.name for f in files]
    fmt         = target_fmt.lower().split("(")[-1].strip(")").strip(".")
    outs, log   = batch_convert(paths, fmt, use_groq)
    return outs or [], log


# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 7 — GRADIO UI + LAUNCH
# ═══════════════════════════════════════════════════════════════════════════

# ── Custom CSS (Dark Navy / Gold theme) ───────────────────────────────────
CUSTOM_CSS = """
@import url('https://fonts.googleapis.com/css2?family=Syne:wght@400;600;700;800&family=DM+Sans:ital,wght@0,300;0,400;0,500;1,300&display=swap');

:root {
    --navy:   #0d1b2a;
    --navy2:  #1b2e45;
    --navy3:  #243b55;
    --gold:   #f0c040;
    --gold2:  #e8a800;
    --silver: #c8d8e8;
    --white:  #f4f8ff;
    --green:  #2ecc71;
    --radius: 12px;
}

body, .gradio-container {
    background: var(--navy) !important;
    font-family: 'DM Sans', sans-serif !important;
    color: var(--white) !important;
}

.app-header {
    background: linear-gradient(135deg, var(--navy2) 0%, var(--navy3) 50%, #1a3a5c 100%);
    border: 1px solid rgba(240,192,64,0.25);
    border-radius: var(--radius);
    padding: 28px 36px;
    margin-bottom: 20px;
    position: relative;
    overflow: hidden;
}
.app-header::before {
    content: '';
    position: absolute; top: 0; left: 0; right: 0; height: 3px;
    background: linear-gradient(90deg, var(--gold2), var(--gold), var(--gold2));
}
.app-header h1 {
    font-family: 'Syne', sans-serif !important;
    font-weight: 800 !important;
    font-size: 2.2rem !important;
    color: var(--white) !important;
    margin: 0 0 6px 0 !important;
}
.app-header h1 span { color: var(--gold); }
.app-header p {
    color: var(--silver) !important;
    font-size: 1rem !important;
    margin: 0 !important;
    font-weight: 300;
}
.badge {
    display: inline-block;
    background: rgba(240,192,64,0.15);
    border: 1px solid var(--gold);
    color: var(--gold);
    font-size: 0.72rem;
    font-weight: 600;
    padding: 2px 10px;
    border-radius: 20px;
    margin-right: 6px;
    font-family: 'Syne', sans-serif;
    letter-spacing: 0.5px;
    text-transform: uppercase;
}
.tab-nav {
    background: var(--navy2) !important;
    border-radius: var(--radius) !important;
    padding: 6px !important;
    border: 1px solid rgba(255,255,255,0.08) !important;
}
.tab-nav button {
    background: transparent !important;
    color: var(--silver) !important;
    border: none !important;
    border-radius: 8px !important;
    font-family: 'Syne', sans-serif !important;
    font-weight: 600 !important;
    font-size: 0.82rem !important;
    padding: 8px 14px !important;
    transition: all 0.2s ease !important;
}
.tab-nav button.selected, .tab-nav button:hover {
    background: var(--gold) !important;
    color: var(--navy) !important;
}
.gradio-group, .gr-group {
    background: var(--navy2) !important;
    border: 1px solid rgba(255,255,255,0.08) !important;
    border-radius: var(--radius) !important;
    padding: 20px !important;
}
button.primary {
    background: linear-gradient(135deg, var(--gold2), var(--gold)) !important;
    color: var(--navy) !important;
    border: none !important;
    border-radius: 8px !important;
    font-family: 'Syne', sans-serif !important;
    font-weight: 700 !important;
    padding: 12px 28px !important;
    transition: transform 0.15s, box-shadow 0.15s !important;
    box-shadow: 0 4px 14px rgba(240,192,64,0.3) !important;
}
button.primary:hover {
    transform: translateY(-2px) !important;
    box-shadow: 0 7px 20px rgba(240,192,64,0.45) !important;
}
input, textarea, select {
    background: rgba(255,255,255,0.05) !important;
    border: 1px solid rgba(255,255,255,0.12) !important;
    color: var(--white) !important;
    border-radius: 8px !important;
}
input:focus, textarea:focus {
    border-color: var(--gold) !important;
    box-shadow: 0 0 0 3px rgba(240,192,64,0.15) !important;
    outline: none !important;
}
label {
    color: var(--silver) !important;
    font-size: 0.85rem !important;
    font-weight: 500 !important;
}
.status-box textarea {
    background: rgba(0,0,0,0.3) !important;
    border: 1px solid rgba(255,255,255,0.1) !important;
    border-radius: 8px !important;
    font-size: 0.82rem !important;
    color: var(--green) !important;
    line-height: 1.6 !important;
}
.section-title {
    font-family: 'Syne', sans-serif;
    font-weight: 700;
    font-size: 1.05rem;
    color: var(--gold);
    padding: 10px 0 8px 0;
    border-bottom: 1px solid rgba(240,192,64,0.2);
    margin-bottom: 14px;
}
.footer-note {
    text-align: center;
    color: rgba(200,216,232,0.45);
    font-size: 0.78rem;
    padding: 16px;
}
::-webkit-scrollbar { width: 6px; height: 6px; }
::-webkit-scrollbar-track { background: var(--navy); }
::-webkit-scrollbar-thumb { background: var(--navy3); border-radius: 4px; }
::-webkit-scrollbar-thumb:hover { background: var(--gold2); }
"""

# ── Build the Gradio app ───────────────────────────────────────────────────
groq_status = (
    "🟢 Groq AI Active — llama-3.3-70b-versatile"
    if GROQ_AVAILABLE
    else "🟡 Groq Offline — Standard Mode Only"
)

with gr.Blocks(css=CUSTOM_CSS, title="AI File Converter Pro", theme=gr.themes.Base()) as app:

    # Header banner
    gr.HTML(f"""
    <div class="app-header">
        <h1>🔄 <span>AI</span> File Converter Pro</h1>
        <p>Enterprise-grade document conversion powered by Groq LLaMA</p>
        <div style="margin-top:12px">
            <span class="badge">PDF</span>
            <span class="badge">Word</span>
            <span class="badge">Excel</span>
            <span class="badge">CSV</span>
            <span class="badge">OCR</span>
            <span class="badge">AI Extract</span>
            <span style="font-size:0.82rem;color:#c8d8e8;margin-left:10px;">{groq_status}</span>
        </div>
    </div>
    """)

    with gr.Tabs():

        # ── Tab 1: PDF → Word ──────────────────────────────────────────────
        with gr.Tab("📄 PDF → Word"):
            gr.HTML('<div class="section-title">📄 PDF to Word Converter</div>')
            with gr.Row():
                with gr.Column(scale=1):
                    p2w_file   = gr.File(label="Upload PDF", file_types=[".pdf"])
                    p2w_groq   = gr.Checkbox(
                        label="🤖 Use Groq AI Vision (best for scanned/image PDFs)",
                        value=False
                    )
                    p2w_btn    = gr.Button("Convert to Word", variant="primary")
                with gr.Column(scale=1):
                    p2w_out    = gr.File(label="⬇️ Download Word Document")
                    p2w_status = gr.Textbox(label="Status", interactive=False, elem_classes=["status-box"])
            with gr.Accordion("ℹ️ Tips", open=False):
                gr.Markdown("""
- **Standard mode** — fast, works for text-based PDFs
- **Groq AI mode** — handles scanned/image PDFs via Vision OCR
- Multi-page PDFs are processed page by page
""")
            p2w_btn.click(handle_pdf_to_word,
                          inputs=[p2w_file, p2w_groq],
                          outputs=[p2w_out, p2w_status])

        # ── Tab 2: Word → PDF ──────────────────────────────────────────────
        with gr.Tab("📝 Word → PDF"):
            gr.HTML('<div class="section-title">📝 Word to PDF Converter</div>')
            with gr.Row():
                with gr.Column(scale=1):
                    w2p_file   = gr.File(label="Upload Word Document (.docx)", file_types=[".docx", ".doc"])
                    w2p_btn    = gr.Button("Convert to PDF", variant="primary")
                with gr.Column(scale=1):
                    w2p_out    = gr.File(label="⬇️ Download PDF")
                    w2p_status = gr.Textbox(label="Status", interactive=False, elem_classes=["status-box"])
            with gr.Accordion("ℹ️ Tips", open=False):
                gr.Markdown("""
- Uses **LibreOffice headless** for high-fidelity conversion
- Fonts, tables, images, and formatting are preserved
- Falls back to docx2pdf if LibreOffice is unavailable
""")
            w2p_btn.click(handle_word_to_pdf,
                          inputs=[w2p_file],
                          outputs=[w2p_out, w2p_status])

        # ── Tab 3: Word → Excel ────────────────────────────────────────────
        with gr.Tab("📊 Word → Excel"):
            gr.HTML('<div class="section-title">📊 Word to Excel — AI Data Extraction</div>')
            with gr.Row():
                with gr.Column(scale=1):
                    w2e_file   = gr.File(label="Upload Word Document (.docx)", file_types=[".docx"])
                    w2e_groq   = gr.Checkbox(
                        label="🤖 Use Groq AI to extract implicit tables (grades, invoices, lists…)",
                        value=True
                    )
                    w2e_btn    = gr.Button("Extract to Excel", variant="primary")
                with gr.Column(scale=1):
                    w2e_out    = gr.File(label="⬇️ Download Excel Workbook")
                    w2e_status = gr.Textbox(label="Status", interactive=False, elem_classes=["status-box"])
            with gr.Accordion("ℹ️ Tips", open=False):
                gr.Markdown("""
- Native Word tables are always extracted
- Groq AI also parses narrative text into structured rows
- Each table becomes a separate Excel sheet
""")
            w2e_btn.click(handle_word_to_excel,
                          inputs=[w2e_file, w2e_groq],
                          outputs=[w2e_out, w2e_status])

        # ── Tab 4: Excel → Word ────────────────────────────────────────────
        with gr.Tab("📋 Excel → Word"):
            gr.HTML('<div class="section-title">📋 Excel to Word — Table Report Generator</div>')
            with gr.Row():
                with gr.Column(scale=1):
                    e2w_file   = gr.File(label="Upload Excel Workbook (.xlsx)", file_types=[".xlsx", ".xls"])
                    e2w_btn    = gr.Button("Generate Word Report", variant="primary")
                with gr.Column(scale=1):
                    e2w_out    = gr.File(label="⬇️ Download Word Report")
                    e2w_status = gr.Textbox(label="Status", interactive=False, elem_classes=["status-box"])
            with gr.Accordion("ℹ️ Tips", open=False):
                gr.Markdown("""
- Each Excel sheet becomes a formatted Word table
- Auto-generates title, date, and page structure
- Header rows use navy background with white text
""")
            e2w_btn.click(handle_excel_to_word,
                          inputs=[e2w_file],
                          outputs=[e2w_out, e2w_status])

        # ── Tab 5: CSV ↔ Excel ─────────────────────────────────────────────
        with gr.Tab("📁 CSV ↔ Excel"):
            gr.HTML('<div class="section-title">📁 CSV ↔ Excel Converter</div>')
            with gr.Row():
                with gr.Column():
                    gr.Markdown("### CSV → Excel")
                    c2e_file   = gr.File(label="Upload CSV File", file_types=[".csv", ".tsv"])
                    c2e_delim  = gr.Dropdown(
                        choices=["Auto-detect", "Comma (,)", "Semicolon (;)", "Tab (\\t)", "Pipe (|)"],
                        value="Auto-detect",
                        label="Delimiter"
                    )
                    c2e_btn    = gr.Button("Convert to Excel", variant="primary")
                    c2e_out    = gr.File(label="⬇️ Download Excel")
                    c2e_status = gr.Textbox(label="Status", interactive=False, elem_classes=["status-box"])
                with gr.Column():
                    gr.Markdown("### Excel → CSV")
                    e2c_file   = gr.File(label="Upload Excel File", file_types=[".xlsx", ".xls"])
                    e2c_delim  = gr.Dropdown(
                        choices=["Comma (,)", "Semicolon (;)", "Tab (\\t)", "Pipe (|)"],
                        value="Comma (,)",
                        label="Output Delimiter"
                    )
                    e2c_btn    = gr.Button("Convert to CSV", variant="primary")
                    e2c_out    = gr.File(label="⬇️ Download CSV")
                    e2c_status = gr.Textbox(label="Status", interactive=False, elem_classes=["status-box"])
            c2e_btn.click(handle_csv_to_excel,
                          inputs=[c2e_file, c2e_delim],
                          outputs=[c2e_out, c2e_status])
            e2c_btn.click(handle_excel_to_csv,
                          inputs=[e2c_file, e2c_delim],
                          outputs=[e2c_out, e2c_status])

        # ── Tab 6: AI Summarizer ───────────────────────────────────────────
        with gr.Tab("🤖 AI Summarizer"):
            gr.HTML('<div class="section-title">🤖 Groq AI Document Summarizer</div>')
            with gr.Row():
                with gr.Column(scale=1):
                    sum_file  = gr.File(
                        label="Upload Document",
                        file_types=[".docx", ".txt", ".csv", ".xlsx"]
                    )
                    sum_style = gr.Radio(
                        choices=["Concise", "Detailed", "Bullet"],
                        value="Concise",
                        label="Summary Style"
                    )
                    sum_btn   = gr.Button("✨ Summarize with Groq AI", variant="primary")
                with gr.Column(scale=1):
                    sum_out   = gr.Textbox(label="AI Summary", lines=16, interactive=False)
            if not GROQ_AVAILABLE:
                gr.HTML(
                    '<div style="color:#f0c040;padding:10px;background:rgba(240,192,64,0.08);'
                    'border-radius:8px">⚠️ Set GROQ_API_KEY environment variable to enable AI.</div>'
                )
            sum_btn.click(handle_summarize, inputs=[sum_file, sum_style], outputs=[sum_out])

        # ── Tab 7: Batch Convert ───────────────────────────────────────────
        with gr.Tab("⚡ Batch Convert"):
            gr.HTML('<div class="section-title">⚡ Batch File Converter</div>')
            with gr.Row():
                with gr.Column(scale=1):
                    batch_files = gr.File(label="Upload Multiple Files", file_count="multiple")
                    batch_fmt   = gr.Dropdown(
                        choices=["Word (.docx)", "PDF (.pdf)", "Excel (.xlsx)", "CSV (.csv)"],
                        value="Word (.docx)",
                        label="Convert ALL files to:"
                    )
                    batch_groq  = gr.Checkbox(label="🤖 Enable Groq AI Enhancement", value=False)
                    batch_btn   = gr.Button("🚀 Start Batch Conversion", variant="primary")
                with gr.Column(scale=1):
                    batch_out = gr.File(label="⬇️ Download Converted Files", file_count="multiple")
                    batch_log = gr.Textbox(
                        label="Conversion Log", lines=14, interactive=False,
                        elem_classes=["status-box"]
                    )
            batch_btn.click(handle_batch,
                            inputs=[batch_files, batch_fmt, batch_groq],
                            outputs=[batch_out, batch_log])

    # Footer
    gr.HTML("""
    <div class="footer-note">
        AI File Converter Pro &nbsp;·&nbsp; Powered by Groq LLaMA-3.3-70b &nbsp;·&nbsp;
        pdf2docx · python-docx · openpyxl · pandas · Gradio
    </div>
    """)


# ═══════════════════════════════════════════════════════════════════════════
#  SECTION 8 — ENTRY POINT
# ═══════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("=" * 60)
    print("  🔄 AI File Converter Pro")
    print("=" * 60)
    print(f"  Groq AI  : {'Active ✅' if GROQ_AVAILABLE else 'Offline ⚠️ (set GROQ_API_KEY)'}")
    print(f"  pdf2docx : {'✅' if PDF2DOCX_AVAILABLE else '❌ not installed'}")
    print(f"  pdf2image: {'✅' if PDF2IMAGE_AVAILABLE else '❌ not installed'}")
    print(f"  Tesseract: {'✅' if TESSERACT_AVAILABLE else '❌ not installed'}")
    print("=" * 60)
    print("  Open: http://localhost:7860")
    print("=" * 60)

    app.launch(
        server_name="0.0.0.0",   # accessible on network
        server_port=7860,
        share=False,             # set True for public Gradio link
        show_error=True,
        debug=False
    )